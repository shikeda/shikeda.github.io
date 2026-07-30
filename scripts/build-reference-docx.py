#!/usr/bin/env python3
"""Build publication/reference.docx, the Pandoc reference document that styles
every Word file generated from the KRM Documentation.

Design brief (see project discussion): combine the academic seriousness of the
TEI Guidelines with the readability of Read the Docs. Concretely: a plain,
undecorated title page; generous margins; a running header showing the book
title on both verso and recto pages (ideally the recto side would show the
current chapter title via a STYLEREF field, but that is disabled for now --
see the comment at odd_header below); a footer with version / page number /
date; quiet sans-serif headings over serif body text; and shaded, monospace
source code.

This script starts from Pandoc's own default reference.docx (so every style
Pandoc's docx writer expects already exists with the right id/behavior/
relationships) and then reworks fonts, sizes, colors, spacing, borders,
shading, page setup, and headers/footers on top of it. Regenerate with:

    python3 scripts/build-reference-docx.py

Requires the `python-docx` package and the `pandoc` executable.
"""

from __future__ import annotations

import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt, RGBColor, Twips

REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = REPO_ROOT / "publication" / "reference.docx"

# ---------------------------------------------------------------------------
# Design tokens
# ---------------------------------------------------------------------------

FONT_HEADING = "Calibri"
FONT_HEADING_EASTASIA = "Yu Gothic"
FONT_BODY = "Cambria"
FONT_BODY_EASTASIA = "Yu Mincho"
FONT_CODE = "Consolas"
FONT_CODE_EASTASIA = "MS Gothic"

INK = RGBColor(0x1A, 0x1A, 0x1A)  # near-black, used for headings
BODY_COLOR = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x59, 0x59, 0x59)  # captions, header/footer, subtitle
RULE_COLOR = "BFBFBF"  # heading/table rules
FAINT_RULE_COLOR = "D9D9D9"  # header/footer rule
CODE_SHADING = "F2F2F2"
TABLE_HEADER_SHADING = "F2F2F2"
HYPERLINK_COLOR = RGBColor(0x05, 0x63, 0xC1)

PAGE_WIDTH = Mm(210)
PAGE_HEIGHT = Mm(297)
MARGIN_SIDE = Mm(30)
MARGIN_TOP = Mm(25)
MARGIN_BOTTOM = Mm(25)
HEADER_DISTANCE = Mm(12.5)
FOOTER_DISTANCE = Mm(12.5)

DOC_TITLE = "KRM Documentation"
DOC_VERSION = "Version 0.9"
DOC_DATE = "July 2026"


# ---------------------------------------------------------------------------
# Low-level OOXML helpers (python-docx has no high-level API for these)
# ---------------------------------------------------------------------------

def set_rpr_fonts(rpr, ascii_font, eastasia_font, hansi_font=None):
    """Set ascii/hAnsi/eastAsia/cs fonts on an existing <w:rPr>."""
    hansi_font = hansi_font or ascii_font
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rStyle = rpr.find(qn("w:rStyle"))
        if rStyle is not None:
            rStyle.addnext(rFonts)
        else:
            rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), hansi_font)
    rFonts.set(qn("w:eastAsia"), eastasia_font)
    rFonts.set(qn("w:cs"), ascii_font)


def style_rpr(style):
    """Return (creating if needed) the <w:rPr> of a style's <w:style> element."""
    el = style.element
    rpr = el.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        el.append(rpr)
    return rpr


def style_ppr(style):
    el = style.element
    ppr = el.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        # <w:pPr> must precede <w:rPr> in a <w:style> element
        rpr = el.find(qn("w:rPr"))
        if rpr is not None:
            rpr.addprevious(ppr)
        else:
            el.append(ppr)
    return ppr


# CT_PPrBase's required child sequence (ECMA-376). python-docx's own
# paragraph-format setters (space_before, tab_stops, ...) already insert at
# the correct position via this same ordering; the pBdr/shd helpers below
# manipulate the OOXML directly and must replicate it by hand, or a strict
# consumer could treat the resulting part as invalid.
_PPR_CHILD_ORDER = (
    "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore", "w:framePr",
    "w:widowControl", "w:numPr", "w:suppressLineNumbers", "w:pBdr", "w:shd",
    "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
    "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
    "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
    "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
)


def insert_ppr_child(ppr, new_el, tag):
    """Insert `new_el` (with qualified name `tag`, e.g. "w:pBdr") into `ppr`
    immediately before its first existing successor, per CT_PPrBase's
    required child order."""
    idx = _PPR_CHILD_ORDER.index(tag)
    for successor_tag in _PPR_CHILD_ORDER[idx + 1:]:
        successor = ppr.find(qn(successor_tag))
        if successor is not None:
            successor.addprevious(new_el)
            return
    ppr.append(new_el)


def set_font(style, *, ascii_font, eastasia_font, size=None, bold=None,
             italic=None, color=None, small_caps=None):
    font = style.font
    if size is not None:
        font.size = size
        # font.size only sets w:sz (ascii/eastAsia size); keep w:szCs (complex
        # script size, e.g. Arabic/Hebrew) in sync so no stale value survives
        # from the base template.
        rpr = style_rpr(style)
        szCs = rpr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rpr.append(szCs)
        szCs.set(qn("w:val"), str(int(size.pt * 2)))
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if color is not None:
        font.color.rgb = color
    if small_caps is not None:
        font.small_caps = small_caps
    set_rpr_fonts(style_rpr(style), ascii_font, eastasia_font)


def set_spacing(style, *, before=None, after=None, line=None,
                 line_rule=None, keep_next=None, page_break_before=None):
    pf = style.paragraph_format
    if before is not None:
        pf.space_before = before
    if after is not None:
        pf.space_after = after
    if line is not None:
        pf.line_spacing = line
    if keep_next is not None:
        pf.keep_with_next = keep_next
    if page_break_before is not None:
        pf.page_break_before = page_break_before


def _get_or_insert_pbdr(ppr):
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        insert_ppr_child(ppr, pbdr, "w:pBdr")
    return pbdr


def add_bottom_border(style, *, size=6, color=RULE_COLOR, space=4):
    pbdr = _get_or_insert_pbdr(style_ppr(style))
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_top_border(style, *, size=4, color=FAINT_RULE_COLOR, space=4):
    pbdr = _get_or_insert_pbdr(style_ppr(style))
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), str(space))
    top.set(qn("w:color"), color)
    pbdr.append(top)


def set_paragraph_shading(style, fill):
    ppr = style_ppr(style)
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        insert_ppr_child(ppr, shd, "w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


# CT_RPr's required child sequence (ECMA-376), abbreviated to the elements
# this script ever touches or might find pre-existing in the base template.
_RPR_CHILD_ORDER = (
    "w:rStyle", "w:rFonts", "w:b", "w:bCs", "w:i", "w:iCs", "w:caps",
    "w:smallCaps", "w:strike", "w:dstrike", "w:outline", "w:shadow",
    "w:emboss", "w:imprint", "w:noProof", "w:snapToGrid", "w:vanish",
    "w:webHidden", "w:color", "w:spacing", "w:w", "w:kern", "w:position",
    "w:sz", "w:szCs", "w:highlight", "w:u", "w:effect", "w:bdr", "w:shd",
    "w:fitText", "w:vertAlign", "w:rtl", "w:cs", "w:em", "w:lang",
    "w:eastAsianLayout", "w:specVanish", "w:oMath",
)


def set_char_shading(style, fill):
    rpr = style_rpr(style)
    shd = rpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        idx = _RPR_CHILD_ORDER.index("w:shd")
        inserted = False
        for successor_tag in _RPR_CHILD_ORDER[idx + 1:]:
            successor = rpr.find(qn(successor_tag))
            if successor is not None:
                successor.addprevious(shd)
                inserted = True
                break
        if not inserted:
            rpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def add_tab_stops(style, stops):
    """stops: list of (position_in_twips, WD_TAB_ALIGNMENT, WD_TAB_LEADER_or_None)."""
    style.paragraph_format.tab_stops.clear_all()
    for pos, align, leader in stops:
        style.paragraph_format.tab_stops.add_tab_stop(
            Twips(pos), align, leader if leader is not None else WD_TAB_LEADER.SPACES
        )


def add_field(paragraph, instr_text, cached_text, *, bold=None, italic=None):
    """Append a legacy w:fldSimple field (auto-updating in Word) holding
    `cached_text` as its placeholder/cached result."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr_text)
    r = OxmlElement("w:r")
    if bold or italic:
        rpr = OxmlElement("w:rPr")
        if bold:
            rpr.append(OxmlElement("w:b"))
        if italic:
            rpr.append(OxmlElement("w:i"))
        r.append(rpr)
    t = OxmlElement("w:t")
    t.text = cached_text
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def add_table_style_pr(table_style, type_, *, bold=None, shading_fill=None):
    """Add or extend the <w:tblStylePr> conditional-formatting block for
    `type_` (e.g. "firstRow"). Reuses an existing block of the same type
    instead of appending a duplicate sibling, since most renderers only
    honor one block per type and Pandoc's base reference.docx already
    ships a "firstRow" block of its own."""
    el = table_style.element
    pr = None
    for existing in el.findall(qn("w:tblStylePr")):
        if existing.get(qn("w:type")) == type_:
            pr = existing
            break
    if pr is None:
        pr = OxmlElement("w:tblStylePr")
        pr.set(qn("w:type"), type_)
        el.append(pr)
    if bold:
        rpr = pr.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            pr.insert(0, rpr)
        rpr.append(OxmlElement("w:b"))
    if shading_fill:
        tcpr = pr.find(qn("w:tcPr"))
        if tcpr is None:
            tcpr = OxmlElement("w:tcPr")
            pr.append(tcpr)
        shd = tcpr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tcpr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), shading_fill)


def set_table_borders(table_style, *, size=4, color=RULE_COLOR):
    el = table_style.element
    tblPr = el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        el.append(tblPr)
    borders = OxmlElement("w:tblBorders")
    for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{tag}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)
        borders.append(edge)
    tblPr.append(borders)
    cellmar = tblPr.find(qn("w:tblCellMar"))
    if cellmar is not None:
        for tag, val in (("top", 60), ("bottom", 60), ("left", 120), ("right", 120)):
            el2 = cellmar.find(qn(f"w:{tag}"))
            if el2 is not None:
                el2.set(qn("w:w"), str(val))


# ---------------------------------------------------------------------------
# Style helpers that operate at the Document level
# ---------------------------------------------------------------------------

def get_style(doc, style_id):
    """Look up a style by its style_id (not its display name), without
    triggering python-docx's by-name-preferred deprecation warning."""
    for style in doc.styles:
        if style.style_id == style_id:
            return style
    return None


def get_or_add_style(doc, style_id, name, style_type, based_on=None):
    style = get_style(doc, style_id)
    if style is not None:
        return style
    style = doc.styles.add_style(name, style_type)
    style.style_id = style_id
    if based_on is not None:
        style.base_style = get_style(doc, based_on)
    return style


def build_toc_style(doc, level, size, *, bold=False, italic=False, indent_mm,
                     usable_width_twips):
    style_id = f"TOC{level}"
    style = get_or_add_style(
        doc, style_id, f"TOC {level}", WD_STYLE_TYPE.PARAGRAPH, based_on="Normal"
    )
    style.quick_style = True
    set_font(
        style,
        ascii_font=FONT_HEADING,
        eastasia_font=FONT_HEADING_EASTASIA,
        size=size,
        bold=bold,
        italic=italic,
        color=INK if level == 1 else BODY_COLOR,
    )
    set_spacing(style, before=Pt(2), after=Pt(4), line=1.15)
    pf = style.paragraph_format
    pf.left_indent = Mm(indent_mm)
    add_tab_stops(
        style,
        [(usable_width_twips, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)],
    )
    return style


# ---------------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------------

def main():
    with tempfile.TemporaryDirectory() as tmp:
        base_path = Path(tmp) / "base-reference.docx"
        subprocess.run(
            ["pandoc", "-o", str(base_path), "--print-default-data-file", "reference.docx"],
            check=True,
        )
        doc = Document(str(base_path))

    # --- Page setup -------------------------------------------------------
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_SIDE
    section.right_margin = MARGIN_SIDE
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE
    section.gutter = 0
    section.different_first_page_header_footer = True  # blank header/footer on the cover page
    doc.settings.odd_and_even_pages_header_footer = True

    usable_width_twips = Emu(PAGE_WIDTH - 2 * MARGIN_SIDE).twips

    # --- Document defaults --------------------------------------------------
    normal = get_style(doc, "Normal")
    set_font(normal, ascii_font=FONT_BODY, eastasia_font=FONT_BODY_EASTASIA,
             size=Pt(11), color=BODY_COLOR)
    set_spacing(normal, after=Pt(8), line=1.2)

    # --- Body text styles ----------------------------------------------------
    body_text = get_style(doc, "BodyText")
    set_font(body_text, ascii_font=FONT_BODY, eastasia_font=FONT_BODY_EASTASIA,
              size=Pt(11), color=BODY_COLOR)
    set_spacing(body_text, before=Pt(0), after=Pt(8), line=1.25)

    first_paragraph = get_style(doc, "FirstParagraph")
    set_font(first_paragraph, ascii_font=FONT_BODY, eastasia_font=FONT_BODY_EASTASIA,
              size=Pt(11), color=BODY_COLOR)
    set_spacing(first_paragraph, before=Pt(0), after=Pt(8), line=1.25)

    compact = get_style(doc, "Compact")
    set_font(compact, ascii_font=FONT_BODY, eastasia_font=FONT_BODY_EASTASIA,
              size=Pt(10.5), color=BODY_COLOR)
    set_spacing(compact, before=Pt(2), after=Pt(2), line=1.2)

    body_text_char = get_style(doc, "BodyTextChar")
    set_rpr_fonts(style_rpr(body_text_char), FONT_BODY, FONT_BODY_EASTASIA)

    block_text = get_style(doc, "BlockText")
    set_font(block_text, ascii_font=FONT_BODY, eastasia_font=FONT_BODY_EASTASIA,
              size=Pt(10.5), italic=True, color=MUTED)
    pf = block_text.paragraph_format
    pf.left_indent = Mm(8)
    pf.right_indent = Mm(8)

    # --- Title page -----------------------------------------------------
    title = get_style(doc, "Title")
    set_font(title, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(32), bold=True, color=INK)
    set_spacing(title, before=Pt(0), after=Pt(18), line=1.1)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = get_style(doc, "Subtitle")
    set_font(subtitle, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(16), bold=False, color=MUTED)
    set_spacing(subtitle, before=Pt(0), after=Pt(36), line=1.2)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    author = get_style(doc, "Author")
    set_font(author, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(13), bold=False, color=INK)
    set_spacing(author, before=Pt(0), after=Pt(4), line=1.2)
    author.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_style = get_style(doc, "Date")
    set_font(date_style, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(11), italic=True, color=MUTED)
    set_spacing(date_style, before=Pt(0), after=Pt(4), line=1.2)
    date_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    abstract_title = get_style(doc, "AbstractTitle")
    set_font(abstract_title, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(12), bold=True, color=INK)
    abstract_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    abstract = get_style(doc, "Abstract")
    set_font(abstract, ascii_font=FONT_BODY, eastasia_font=FONT_BODY_EASTASIA,
              size=Pt(10.5), italic=True, color=MUTED)
    abstract.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Headings ----------------------------------------------------------
    h1 = get_style(doc, "Heading1")
    set_font(h1, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(16), bold=True, color=INK)
    set_spacing(h1, before=Pt(0), after=Pt(14), line=1.1,
                keep_next=True, page_break_before=True)
    add_bottom_border(h1, size=6, color=RULE_COLOR, space=6)

    h2 = get_style(doc, "Heading2")
    set_font(h2, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(14), bold=True, color=INK)
    set_spacing(h2, before=Pt(20), after=Pt(8), line=1.1, keep_next=True)

    h3 = get_style(doc, "Heading3")
    set_font(h3, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(12), bold=True, color=INK)
    set_spacing(h3, before=Pt(14), after=Pt(6), line=1.1, keep_next=True)

    h4 = get_style(doc, "Heading4")
    set_font(h4, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(11), bold=True, italic=True, color=INK)
    set_spacing(h4, before=Pt(10), after=Pt(4), line=1.1, keep_next=True)

    # Headings 5-9 exist for schema completeness; style them as a quiet,
    # smaller continuation of the H4 look so a deeply nested page never
    # falls back to Word's raw defaults.
    for n in range(5, 10):
        h = get_style(doc, f"Heading{n}")
        set_font(h, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
                  size=Pt(10.5), bold=True, italic=True, color=MUTED)
        set_spacing(h, before=Pt(8), after=Pt(4), line=1.1, keep_next=True)

    section_number = get_style(doc, "SectionNumber")
    set_rpr_fonts(style_rpr(section_number), FONT_HEADING, FONT_HEADING_EASTASIA)
    section_number.font.color.rgb = INK
    section_number.font.bold = True

    # --- Source code ---------------------------------------------------
    source_code = get_or_add_style(
        doc, "SourceCode", "Source Code", WD_STYLE_TYPE.PARAGRAPH, based_on="Normal"
    )
    set_font(source_code, ascii_font=FONT_CODE, eastasia_font=FONT_CODE_EASTASIA,
              size=Pt(9.5), color=RGBColor(0x2B, 0x2B, 0x2B))
    set_spacing(source_code, before=Pt(6), after=Pt(6), line=1.25)
    pf = source_code.paragraph_format
    pf.left_indent = Mm(4)
    pf.right_indent = Mm(4)
    set_paragraph_shading(source_code, CODE_SHADING)

    verbatim_char = get_style(doc, "VerbatimChar")
    set_rpr_fonts(style_rpr(verbatim_char), FONT_CODE, FONT_CODE_EASTASIA)
    verbatim_char.font.size = Pt(10)
    verbatim_char.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)
    set_char_shading(verbatim_char, CODE_SHADING)

    # --- Captions / figures ------------------------------------------------
    for style_id in ("Caption", "TableCaption", "ImageCaption"):
        st = get_style(doc, style_id)
        set_font(st, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
                  size=Pt(9.5), italic=True, color=MUTED)
        set_spacing(st, before=Pt(4), after=Pt(10), line=1.15)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    figure = get_style(doc, "Figure")
    figure.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(figure, before=Pt(10), after=Pt(2))

    captioned_figure = get_style(doc, "CaptionedFigure")
    captioned_figure.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(captioned_figure, before=Pt(10), after=Pt(2))

    # --- Hyperlinks, footnotes, bibliography --------------------------------
    hyperlink = get_style(doc, "Hyperlink")
    hyperlink.font.color.rgb = HYPERLINK_COLOR
    hyperlink.font.underline = True

    footnote_text = get_style(doc, "FootnoteText")
    set_font(footnote_text, ascii_font=FONT_BODY, eastasia_font=FONT_BODY_EASTASIA,
              size=Pt(9.5), color=BODY_COLOR)

    bibliography = get_style(doc, "Bibliography")
    set_font(bibliography, ascii_font=FONT_BODY, eastasia_font=FONT_BODY_EASTASIA,
              size=Pt(10.5), color=BODY_COLOR)
    pf = bibliography.paragraph_format
    pf.left_indent = Mm(8)
    pf.first_line_indent = Mm(-8)
    set_spacing(bibliography, after=Pt(6), line=1.15)

    definition_term = get_style(doc, "DefinitionTerm")
    set_font(definition_term, ascii_font=FONT_BODY, eastasia_font=FONT_BODY_EASTASIA,
              size=Pt(11), bold=True, color=INK)

    definition = get_style(doc, "Definition")
    set_font(definition, ascii_font=FONT_BODY, eastasia_font=FONT_BODY_EASTASIA,
              size=Pt(11), color=BODY_COLOR)
    definition.paragraph_format.left_indent = Mm(8)

    # --- Table ---------------------------------------------------------
    table_style = get_style(doc, "Table")
    set_font(table_style, ascii_font=FONT_BODY, eastasia_font=FONT_BODY_EASTASIA,
              size=Pt(10), color=BODY_COLOR)
    set_table_borders(table_style, size=4, color=RULE_COLOR)
    add_table_style_pr(table_style, "firstRow", bold=True, shading_fill=TABLE_HEADER_SHADING)

    # --- Table of contents ---------------------------------------------
    toc_heading = get_style(doc, "TOCHeading")
    set_font(toc_heading, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(16), bold=True, color=INK)
    set_spacing(toc_heading, before=Pt(0), after=Pt(18), line=1.1,
                page_break_before=True)
    add_bottom_border(toc_heading, size=6, color=RULE_COLOR, space=6)

    build_toc_style(doc, 1, Pt(12), bold=True, indent_mm=0,
                     usable_width_twips=usable_width_twips)
    build_toc_style(doc, 2, Pt(11), indent_mm=6,
                     usable_width_twips=usable_width_twips)
    build_toc_style(doc, 3, Pt(10.5), indent_mm=12,
                     usable_width_twips=usable_width_twips)
    build_toc_style(doc, 4, Pt(10), italic=True, indent_mm=18,
                     usable_width_twips=usable_width_twips)

    # --- Header / footer text styles ----------------------------------------
    header_style = get_or_add_style(doc, "Header", "Header", WD_STYLE_TYPE.PARAGRAPH,
                                     based_on="Normal")
    set_font(header_style, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(9), color=MUTED)
    set_spacing(header_style, after=Pt(0), line=1.0)
    add_bottom_border(header_style, size=4, color=FAINT_RULE_COLOR, space=6)
    # clear Word's inherited center/right tab stops (Letter-width defaults)
    # before the section's own margins are applied via the header content.
    ppr = style_ppr(header_style)
    tabs = ppr.find(qn("w:tabs"))
    if tabs is not None:
        ppr.remove(tabs)

    footer_style = get_or_add_style(doc, "Footer", "Footer", WD_STYLE_TYPE.PARAGRAPH,
                                     based_on="Normal")
    set_font(footer_style, ascii_font=FONT_HEADING, eastasia_font=FONT_HEADING_EASTASIA,
              size=Pt(9), color=MUTED)
    set_spacing(footer_style, before=Pt(0), line=1.0)
    add_top_border(footer_style, size=4, color=FAINT_RULE_COLOR, space=6)
    ppr = style_ppr(footer_style)
    tabs = ppr.find(qn("w:tabs"))
    if tabs is not None:
        ppr.remove(tabs)
    add_tab_stops(
        footer_style,
        [
            (usable_width_twips // 2, WD_TAB_ALIGNMENT.CENTER, None),
            (usable_width_twips, WD_TAB_ALIGNMENT.RIGHT, None),
        ],
    )

    # --- Header/footer content ----------------------------------------------
    # First page (the cover): intentionally blank, per the "no logo, no
    # decoration" brief -- different_first_page_header_footer leaves these
    # empty by default, so nothing further to do here beyond not writing
    # into section.first_page_header / .first_page_footer.

    even_header = section.even_page_header
    even_header.is_linked_to_previous = False
    p = even_header.paragraphs[0]
    p.style = header_style
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(DOC_TITLE)

    odd_header = section.header  # "default" header = odd/recto pages
    odd_header.is_linked_to_previous = False
    p = odd_header.paragraphs[0]
    p.style = header_style
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Ideally this would be a STYLEREF "Heading 1" field, so the recto header
    # shows the current chapter title and updates automatically as the
    # reader scrolls (see add_field() below, still used for the PAGE field
    # in the footer). STYLEREF resolves to the nearest *preceding* Heading 1
    # paragraph, and our front matter (title page, Document Information,
    # Revision History, the Table of Contents field itself) all comes before
    # any Heading 1 -- so on every page up to and including the ToC, Word has
    # no preceding heading to resolve against and renders a field error
    # ("Error! No text of specified style in document.") instead of a blank
    # or the next heading. Disabled for now: falls back to the same static
    # "KRM Documentation" text as the verso header. Revisit once the
    # front-matter/ToC pages either get their own non-STYLEREF header or the
    # book gets a real Heading 1 ahead of the ToC to anchor the field.
    p.add_run(DOC_TITLE)

    even_footer = section.even_page_footer
    even_footer.is_linked_to_previous = False
    p = even_footer.paragraphs[0]
    p.style = footer_style
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(DOC_VERSION)
    p.add_run("\t")
    add_field(p, "PAGE", "1")
    p.add_run("\t" + DOC_DATE)

    odd_footer = section.footer
    odd_footer.is_linked_to_previous = False
    p = odd_footer.paragraphs[0]
    p.style = footer_style
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(DOC_VERSION)
    p.add_run("\t")
    add_field(p, "PAGE", "1")
    p.add_run("\t" + DOC_DATE)

    # --- Save ------------------------------------------------------------
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    sys.exit(main())
